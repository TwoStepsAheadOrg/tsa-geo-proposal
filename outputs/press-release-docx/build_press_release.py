from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "/Users/twostepsahead/Desktop/twostepsahead/tsa-geo-proposal/outputs/press-release-docx/투스텝스어헤드_플루랭크_보도자료_원고.docx"


TITLE = "투스텝스어헤드, AI 답변 ‘인용 가능성’ 예측하는 플루랭크 선보여… 12개국+ AI 검색 데이터 기반 글로벌 GEO 인프라 구축"

ARTICLE_PARAGRAPHS = [
    "AI 마케팅 테크 기업 주식회사 투스텝스어헤드가 생성형 AI 검색 시대에 대응하는 글로벌 GEO 솔루션 ‘플루랭크(Plurank)’를 선보인다고 밝혔다.",
    "플루랭크는 ChatGPT, Gemini, Perplexity, Google AI 등 생성형 AI 기반 답변 환경에서 브랜드가 어떤 질문에 노출되고, 어떤 출처를 근거로 인용되며, 경쟁사 대비 어떤 위치에 있는지를 분석하는 GEO(Generative Engine Optimization) 서비스다. 기존 SEO가 검색 결과 페이지의 순위에 초점을 맞췄다면, 플루랭크는 AI가 답변을 생성할 때 참고하는 질문, 출처, 콘텐츠, 커뮤니티, 소셜 신호까지 함께 분석한다.",
    "특히 플루랭크는 단순한 AI 검색 노출 측정 도구를 넘어, 콘텐츠가 발행되기 전 AI 답변과 외부 채널에서 인용·추천될 가능성을 예측하는 데 초점을 맞췄다. 투스텝스어헤드는 자체 예측 모델 ‘플루오라(Pluora)’를 기반으로 AI 검색, SNS, 커뮤니티, 뉴스·리뷰 매체 등 여러 신호를 통합 분석해 브랜드의 글로벌 발견 가능성을 정량화하고 있다.",
    "최근 검색 시장은 빠르게 변화하고 있다. Gartner는 2026년까지 전통 검색 엔진 사용량이 25% 감소할 수 있다고 전망했으며, Adobe Analytics는 미국 리테일 사이트의 생성형 AI 유입이 크게 증가했다고 발표했다. Semrush 역시 ChatGPT 프롬프트 상당수가 기존 검색 키워드와 일치하지 않는다고 분석했다. 이는 소비자가 더 이상 단순 키워드가 아니라 자연어 질문을 통해 제품과 서비스를 탐색하고 있음을 보여준다.",
    "투스텝스어헤드는 이러한 변화 속에서 플루랭크를 ‘AI 답변 시대의 글로벌 마케팅 인프라’로 포지셔닝하고 있다. 플루랭크는 국가·언어·플랫폼별 질문 데이터를 수집하고, AI 답변 내 브랜드 언급 여부, 인용 URL, 직접·간접 출처, 경쟁사 노출 현황을 테이블과 월간 리포트로 제공한다. 이후 분석 결과를 바탕으로 AI가 인용할 수 있는 FAQ, 비교형 콘텐츠, 리뷰형 콘텐츠, 커뮤니티 답변, SNS·숏폼 소재 등 실행 콘텐츠까지 연결한다.",
    "회사 측은 플루랭크가 12개국 이상의 AI 검색 데이터를 기반으로 글로벌 브랜드와 마케팅 에이전시가 자체적으로 구축하기 어려운 AI 답변 모니터링 체계를 제공한다고 설명했다. 해외 시장에서는 같은 제품이라도 국가별 질문 방식과 인용 출처가 달라지기 때문에, 단순 번역이 아닌 국가별 질문 데이터와 출처 전략이 중요하다는 것이다.",
    "플루랭크의 주요 고객군은 해외 진출을 준비하는 브랜드, 병원·뷰티·관광 등 고관여 업종, 그리고 고객사에 AI 검색 리포트를 제공해야 하는 마케팅 에이전시다. 에이전시에는 고객사별 AI 검색 리포트와 경쟁사 비교 데이터를 제공하고, 브랜드에는 GEO 콘텐츠 제작과 채널 실행까지 포함한 운영 패키지를 제공하는 방식이다.",
    "김건희 투스텝스어헤드 대표는 “AI 검색 시대의 마케팅은 검색 1순위를 만드는 일이 아니라, AI가 답변을 만들 때 어떤 근거로 우리 브랜드를 설명하게 할 것인지의 문제”라며 “플루랭크는 브랜드가 국가별 AI 답변 안에서 어떻게 발견되고 인용되는지를 데이터로 확인하고, 그 결과를 실제 콘텐츠와 매출 기회로 연결하는 것을 목표로 한다”고 말했다.",
    "투스텝스어헤드는 향후 플루랭크를 중심으로 GEO 데이터 리포트, 콘텐츠 실행, 글로벌 채널 전략, 리드 신호 분석을 결합한 AI 마케팅 운영 모델을 확대할 계획이다. 장기적으로는 SaaS 형태의 셀프서비스, API 기반 데이터 제공, AI 에이전트형 마케팅 운영 서비스로 확장한다는 방침이다.",
]

SOURCES = [
    ("Gartner", "https://www.gartner.com/en/newsroom/press-releases/2024-02-19-gartner-predicts-search-engine-volume-will-drop-25-percent-by-2026-due-to-ai-chatbots-and-other-virtual-agents"),
    ("Adobe Analytics", "https://blog.adobe.com/en/publish/2025/03/17/adobe-analytics-traffic-to-us-retail-websites-from-generative-ai-sources-jumps-1200-percent"),
    ("Semrush", "https://www.semrush.com/blog/chatgpt-search-insights/"),
    ("투스텝스어헤드 공식 홈페이지", "https://www.twostepsahead.co.kr/"),
]


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Calibri"
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")
    rfonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
    rfonts.set(qn("w:cs"), "Calibri")


def set_style_font(style, size=None, bold=None, color=None):
    style.font.name = "Calibri"
    if size is not None:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    if color is not None:
        style.font.color.rgb = RGBColor.from_string(color)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")
    rfonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
    rfonts.set(qn("w:cs"), "Calibri")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color="000000"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=9.5, bold=bold, color=color)


def set_cell_width(cell, width_dxa):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = tc_pr.tcW
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))


def set_table_geometry(table, widths):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_pr.append(tbl_layout)

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])


doc = Document()

section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
set_style_font(styles["Normal"], size=11, color="111111")
styles["Normal"].paragraph_format.space_after = Pt(6)
styles["Normal"].paragraph_format.line_spacing = 1.10

for style_name, size, color, before, after in [
    ("Heading 1", 16, "2E74B5", 16, 8),
    ("Heading 2", 13, "2E74B5", 12, 6),
    ("Heading 3", 12, "1F4D78", 8, 4),
]:
    set_style_font(styles[style_name], size=size, bold=True, color=color)
    styles[style_name].paragraph_format.space_before = Pt(before)
    styles[style_name].paragraph_format.space_after = Pt(after)
    styles[style_name].paragraph_format.line_spacing = 1.10

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer_run = footer.add_run("TwoStepsAhead · Plurank 보도자료 원고")
set_run_font(footer_run, size=8.5, color="777777")

label = doc.add_paragraph()
label.paragraph_format.space_after = Pt(4)
label_run = label.add_run("보도자료 원고")
set_run_font(label_run, size=10, bold=True, color="2E74B5")

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(10)
title_run = title.add_run(TITLE)
set_run_font(title_run, size=18, bold=True, color="0B2545")

meta = doc.add_table(rows=4, cols=2)
meta.style = "Table Grid"
set_table_geometry(meta, [2100, 7260])
meta_rows = [
    ("작성일", "2026년 5월 26일"),
    ("기업명", "주식회사 투스텝스어헤드 (TwoStepsAhead Co., Ltd.)"),
    ("서비스명", "플루랭크(Plurank) · AI 답변 인용 가능성 예측 및 글로벌 GEO 인프라"),
    ("핵심 키워드", "투스텝스어헤드, 플루랭크, Plurank, GEO, 생성형 AI 검색, AI 검색 최적화, 글로벌 마케팅"),
]
for idx, row in enumerate(meta.rows):
    shade_cell(row.cells[0], "F2F4F7")
    set_cell_text(row.cells[0], meta_rows[idx][0], bold=True, color="1F4D78")
    set_cell_text(row.cells[1], meta_rows[idx][1])

doc.add_heading("기사 원고", level=1)
for text in ARTICLE_PARAGRAPHS:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(text)
    set_run_font(run, size=11, color="111111")

doc.add_heading("참고 근거", level=1)
for name, url in SOURCES:
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{name}: {url}")
    set_run_font(run, size=9.5, color="555555")

doc.core_properties.title = TITLE
doc.core_properties.author = "TwoStepsAhead"
doc.core_properties.subject = "Plurank press release draft"
doc.core_properties.keywords = "투스텝스어헤드, 플루랭크, Plurank, GEO, 생성형 AI 검색"
doc.save(OUTPUT)
print(OUTPUT)
